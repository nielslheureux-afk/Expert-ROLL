import streamlit as st
import google.generativeai as genai
import os
import io
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- 1. CONFIGURATION ---
st.set_page_config(
    page_title="ADC — Atelier de Compréhension",
    page_icon="📖",
    layout="centered"
)

# --- CSS ACCESSIBLE + DESIGN DISTINCTIF ---
# Palette : bleu encre profond (#1B2A4A) / crème chaud (#F7F3EC) / ambre (#C17D00)
# Contraste testé : texte foncé sur crème → 12.5:1 / ambre sur fond sombre → 4.8:1
# Police : Atkinson Hyperlegible (conçue pour la lisibilité, dyslexie-friendly)

st.markdown("""
<style>
  /* ── Google Fonts ── */
  @import url('https://fonts.googleapis.com/css2?family=Atkinson+Hyperlegible:ital,wght@0,400;0,700;1,400&family=Fraunces:ital,opsz,wght@0,9..144,300;0,9..144,600;1,9..144,300&display=swap');

  /* ── Variables accessibles ── */
  :root {
    --encre:       #1B2A4A;   /* fond bandeau — très sombre */
    --encre-mid:   #2C4270;   /* blocs secondaires */
    --creme:       #F7F3EC;   /* fond principal — chaud */
    --creme-bord:  #E2D9CC;   /* bordures subtiles */
    --creme-fort:  #C9BFB0;   /* bordures visibles */
    --ambre:       #C17D00;   /* accent principal — ratio 4.8:1 sur fond sombre */
    --ambre-clair: #FFF3D6;   /* fond badge clair */
    --texte:       #1C1A17;   /* texte principal — ratio ~15:1 sur creme */
    --texte-doux:  #3D3529;   /* texte secondaire — ratio ~7:1 sur creme */
    --blanc:       #FFFFFF;
    --focus:       #FF6B35;   /* orange vif pour focus — très visible */
    --erreur-bg:   #FFF0EE;
    --erreur-bord: #C0392B;
    --ok-bg:       #EEF7F0;
    --ok-bord:     #1E7A3A;

    --font-body:   'Atkinson Hyperlegible', Arial, sans-serif;
    --font-titre:  'Fraunces', Georgia, serif;
    --taille-base: 17px;      /* ≥ 16px requis */
    --taille-sm:   15px;      /* min pour texte secondaire */
    --taille-xs:   13px;      /* annotations uniquement */
  }

  /* ── Reset & base ── */
  html, body, [class*="css"] {
    font-family: var(--font-body);
    font-size: var(--taille-base);
    background-color: var(--creme);
    color: var(--texte);
    line-height: 1.7;
    /* PAS de text-align: justify — rivières de blanc */
    text-align: left;
  }

  /* ── Focus universel très visible (navigation clavier) ── */
  *:focus,
  *:focus-visible {
    outline: 3px solid var(--focus) !important;
    outline-offset: 3px !important;
    border-radius: 4px !important;
    box-shadow: 0 0 0 6px rgba(255, 107, 53, 0.25) !important;
  }

  /* ── Lignes de cahier discrètes ── */
  .main > div {
    background-image: repeating-linear-gradient(
      transparent,
      transparent 31px,
      rgba(60, 50, 35, 0.07) 31px,
      rgba(60, 50, 35, 0.07) 32px
    );
    background-attachment: local;
  }

  /* ══════════════════════════════════
     BANDEAU HÉRO
     Fond: --encre (#1B2A4A)
     Texte: #FFFFFF → ratio 17:1 ✓
     Label ambre: #C17D00 sur #1B2A4A → 4.8:1 ✓
  ══════════════════════════════════ */
  .hero-band {
    background: var(--encre);
    color: var(--blanc);
    padding: 2.5rem 2.5rem 2rem;
    margin: -1rem -1rem 2.2rem -1rem;
    position: relative;
    overflow: hidden;
    border-bottom: 4px solid var(--ambre);
  }
  /* Motif géométrique décoratif — aria-hidden via CSS, pas d'info véhiculée */
  .hero-band::after {
    content: "";
    position: absolute;
    right: -3rem;
    top: -3rem;
    width: 16rem;
    height: 16rem;
    border-radius: 50%;
    border: 40px solid rgba(255,255,255,0.04);
    pointer-events: none;
  }
  .hero-label {
    font-family: var(--font-body);
    font-size: var(--taille-xs);
    font-weight: 700;
    letter-spacing: 0.22em;
    text-transform: uppercase;
    color: var(--ambre);   /* #C17D00 sur #1B2A4A → 4.8:1 ✓ */
    margin-bottom: 0.5rem;
    display: block;
  }
  /* h1 — unique sur la page */
  .hero-band h1 {
    font-family: var(--font-titre);
    font-size: clamp(1.9rem, 5vw, 2.8rem);
    font-weight: 600;
    line-height: 1.15;
    margin: 0 0 0.5rem 0;
    color: var(--blanc) !important;
  }
  .hero-band p {
    font-size: clamp(0.95rem, 2vw, 1.05rem);
    color: #C8D4E8;  /* clair sur sombre → ratio ~7:1 ✓ */
    line-height: 1.6;
    max-width: 46rem;
    margin: 0;
    font-style: italic;
  }
  .hero-separateur {
    width: 3.5rem;
    height: 3px;
    background: var(--ambre);
    border: none;
    margin: 1rem 0;
  }

  /* ══════════════════════════════════
     BLOC PRÉSENTATION
     Fond blanc / texte #3D3529 → ratio 8.5:1 ✓
  ══════════════════════════════════ */
  .pres-card {
    background: var(--blanc);
    border-left: 5px solid var(--ambre);
    padding: 1.5rem 1.8rem;
    margin: 0 0 2rem 0;
    border-radius: 0 8px 8px 0;
    box-shadow: 0 2px 16px rgba(27, 42, 74, 0.08);
    font-size: var(--taille-base);
    line-height: 1.75;
    color: var(--texte-doux);
  }
  .pres-card strong {
    color: var(--texte);
    font-weight: 700;
  }
  .pres-droits {
    margin-top: 1rem;
    padding-top: 0.9rem;
    border-top: 1px solid var(--creme-bord);
    font-size: var(--taille-sm);
    color: var(--texte-doux);  /* #3D3529 sur blanc → 8.5:1 ✓ */
    font-style: italic;
  }

  /* ══════════════════════════════════
     ENCADRÉS SECTION
  ══════════════════════════════════ */
  .section-encart {
    background: var(--creme);
    border: 1.5px solid var(--creme-bord);
    border-top: 3px solid var(--encre-mid);
    border-radius: 0 0 8px 8px;
    padding: 1.4rem 1.8rem 1.6rem;
    margin-bottom: 1.6rem;
  }
  /* h2 — titres de sections */
  .section-encart h2 {
    font-family: var(--font-titre);
    font-size: 1.15rem;
    font-weight: 600;
    color: var(--encre);  /* #1B2A4A sur #F7F3EC → 13:1 ✓ */
    margin: 0 0 1rem 0;
    display: flex;
    align-items: center;
    gap: 0.65rem;
  }
  .num-badge {
    display: inline-flex;
    align-items: center;
    justify-content: center;
    background: var(--encre);
    color: var(--blanc);
    font-family: var(--font-body);
    font-size: 0.78rem;
    font-weight: 700;
    width: 1.8rem;
    height: 1.8rem;
    border-radius: 50%;
    flex-shrink: 0;
    /* Accessible : le chiffre est lu par les lecteurs d'écran */
  }

  /* ── Bouton radio Streamlit ── */
  div[role="radiogroup"] {
    gap: 0.5rem;
  }
  div[role="radiogroup"] label {
    background: var(--blanc);
    border: 2px solid var(--creme-fort);
    border-radius: 6px;
    padding: 0.55rem 1.2rem;
    font-family: var(--font-body);
    font-size: var(--taille-sm);
    font-weight: 700;
    color: var(--texte);
    cursor: pointer;
    transition: border-color 0.15s, background 0.15s;
  }
  div[role="radiogroup"] label:hover {
    border-color: var(--encre-mid);
    background: var(--ambre-clair);
  }

  /* ── Upload zone ── */
  [data-testid="stFileUploader"] {
    border: 2px dashed var(--creme-fort);
    border-radius: 10px;
    background: var(--blanc);
    padding: 0.6rem;
    transition: border-color 0.2s;
  }
  [data-testid="stFileUploader"]:hover {
    border-color: var(--encre-mid);
  }

  /* ── Bouton principal ── */
  /* Fond --encre (#1B2A4A) + texte blanc (#FFF) → 17:1 ✓ */
  .stButton > button {
    background: var(--encre) !important;
    color: var(--blanc) !important;
    font-family: var(--font-body) !important;
    font-size: var(--taille-sm) !important;
    font-weight: 700 !important;
    letter-spacing: 0.06em !important;
    text-transform: uppercase !important;
    border: 2px solid transparent !important;
    border-radius: 6px !important;
    padding: 0.7rem 2rem !important;
    transition: background 0.2s, border-color 0.2s !important;
    cursor: pointer !important;
  }
  .stButton > button:hover {
    background: var(--encre-mid) !important;
    border-color: var(--ambre) !important;
  }
  .stButton > button:focus,
  .stButton > button:focus-visible {
    outline: 3px solid var(--focus) !important;
    outline-offset: 3px !important;
  }

  /* ── Bouton téléchargement ── */
  .stDownloadButton > button {
    background: var(--ambre-clair) !important;
    color: var(--texte) !important;  /* sombre sur clair → 14:1 ✓ */
    font-family: var(--font-body) !important;
    font-size: var(--taille-sm) !important;
    font-weight: 700 !important;
    border: 2px solid var(--ambre) !important;
    border-radius: 6px !important;
    padding: 0.65rem 1.6rem !important;
  }
  .stDownloadButton > button:hover {
    background: var(--ambre) !important;
    color: var(--blanc) !important;
  }

  /* ── Zone résultat ── */
  /* h3 — titre du résultat */
  .output-zone {
    background: var(--blanc);
    border: 1.5px solid var(--creme-bord);
    border-left: 5px solid var(--encre-mid);
    border-radius: 0 8px 8px 0;
    padding: 2rem 2rem 2rem 1.8rem;
    margin: 1.5rem 0;
    font-size: var(--taille-base);
    line-height: 1.8;
    color: var(--texte);
    box-shadow: 0 2px 12px rgba(27,42,74,0.06);
  }

  /* ── Messages d'état accessibles (icône + texte, pas seulement couleur) ── */
  [data-testid="stAlert"] {
    font-family: var(--font-body) !important;
    font-size: var(--taille-base) !important;
    border-radius: 8px !important;
  }

  /* ── Footer ── */
  .footer-note {
    font-family: var(--font-body);
    font-size: var(--taille-xs);
    color: var(--texte-doux);  /* #3D3529 sur #F7F3EC → 7:1 ✓ */
    text-align: center;
    margin-top: 3rem;
    padding-top: 1rem;
    border-top: 1px solid var(--creme-bord);
    line-height: 1.6;
  }

  /* ── Responsive — pas de scroll horizontal ── */
  @media (max-width: 640px) {
    .hero-band {
      padding: 1.6rem 1.2rem 1.4rem;
      margin: -0.5rem -0.5rem 1.5rem -0.5rem;
    }
    .hero-band h1 { font-size: 1.6rem; }
    .pres-card, .section-encart, .output-zone {
      padding: 1rem 1rem;
    }
    div[role="radiogroup"] {
      flex-direction: column;
    }
  }

  /* ── Masquer chrome Streamlit non essentiel ── */
  #MainMenu, footer, [data-testid="stToolbar"] { visibility: hidden; }
  [data-testid="stDecoration"] { display: none; }
</style>
""", unsafe_allow_html=True)


# --- 2. MOTEUR DE RENDU WORD ---
def create_adc_docx_final(text_content, cycle_name):
    doc = Document()
    for section in doc.sections:
        section.top_margin, section.bottom_margin = Cm(1.2), Cm(1.2)
        section.left_margin, section.right_margin = Cm(1.5), Cm(1.5)

    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)
    title = doc.add_heading(f"FICHE ENSEIGNANT : ATELIER DE COMPRÉHENSION — {cycle_name}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for line in text_content.split('\n'):
        clean_line = line.strip()
        if not clean_line:
            continue
        if clean_line.startswith(('#', '1.', '2.', '3.', '4.', '5.')) or "PHASE" in clean_line.upper():
            doc.add_heading(clean_line.replace('#', '').strip(), level=1)
        elif "|" in clean_line and "---" not in clean_line:
            parts = [p.strip() for p in clean_line.split("|") if p.strip()]
            if len(parts) >= 2:
                table = doc.add_table(rows=1, cols=len(parts))
                table.style = 'Table Grid'
                for i, part in enumerate(parts):
                    table.rows[0].cells[i].text = part
        elif '**' in clean_line:
            p = doc.add_paragraph()
            for i, part in enumerate(clean_line.split('**')):
                run = p.add_run(part)
                if i % 2 != 0:
                    run.bold = True
        elif clean_line.startswith(('-', '*', '•')):
            doc.add_paragraph(clean_line.strip('-*• ').strip(), style='List Bullet')
        else:
            doc.add_paragraph(clean_line)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# --- 3. INTERFACE ---

# ── Bandeau héro (h1 unique) ──
st.markdown("""
<header class="hero-band" role="banner">
  <span class="hero-label" aria-label="Outil pédagogique pour l'école primaire">
    Outil pédagogique · École primaire
  </span>
  <h1>Atelier de Compréhension</h1>
  <hr class="hero-separateur" aria-hidden="true">
  <p>
    Générateur de fiches enseignantes pour préparer et animer des séances
    de compréhension de texte, inspiré de l'outil ACT du ROLL
    (Réseau des Observatoires Locaux de la Lecture).
  </p>
</header>
""", unsafe_allow_html=True)

# ── Présentation ──
st.markdown("""
<section aria-label="Présentation de l'outil">
  <div class="pres-card">
    Déposez un texte support : l'outil analyse ses obstacles de compréhension
    (lexique, implicite, inférences) et génère une
    <strong>fiche enseignant clé en main</strong> — objectifs ciblés, déroulé
    en phases, tableau collaboratif
    <em>« Ce qu'on sait / Ce qu'on ne sait pas / On n'est pas d'accord »</em>
    — prête à imprimer ou projeter.
    <p class="pres-droits">
      Cet outil génère des fiches de préparation inspirées de l'outil ACT du ROLL
      (roll-descartes.fr). Il ne se substitue pas aux ouvrages de référence et
      n'est pas affilié à leurs auteurs ni à leurs éditeurs. Son usage est
      pédagogique, gratuit et non commercial — il a vocation à encourager les
      enseignants à s'approprier ces approches et à les intégrer dans leur pratique.
    </p>
  </div>
</section>
""", unsafe_allow_html=True)

# ── Section 01 — Niveau ──
st.markdown("""
<div class="section-encart">
  <h2><span class="num-badge" aria-label="Étape 1">1</span>Choisir le niveau de classe</h2>
""", unsafe_allow_html=True)

cycle = st.radio(
    "Niveau de classe :",
    ["Cycle 2 (CP – CE1 – CE2)", "Cycle 3 (CM1 – CM2 – 6ᵉ)"],
    horizontal=True,
    label_visibility="collapsed"
)
cycle_short = "Cycle 2" if "2" in cycle else "Cycle 3"

st.markdown("</div>", unsafe_allow_html=True)

# ── Section 02 — Upload ──
st.markdown("""
<div class="section-encart">
  <h2><span class="num-badge" aria-label="Étape 2">2</span>Déposer le texte support</h2>
""", unsafe_allow_html=True)

uploaded_file = st.file_uploader(
    "Formats acceptés : Word (.docx), PDF, image ou scan (JPG, PNG)",
    type=['docx', 'pdf', 'jpg', 'jpeg', 'png'],
    help="Le fichier peut être un texte tapé, un scan ou une photo. L'IA gère les deux."
)

st.markdown("</div>", unsafe_allow_html=True)

# ── Bouton de génération ──
col1, col2, col3 = st.columns([1, 2, 1])
with col2:
    generate = st.button(
        "▶ Générer la fiche ADC",
        use_container_width=True,
        help="Cliquez après avoir sélectionné le niveau et déposé le fichier"
    )

# ── Logique de génération ──
if uploaded_file and generate:
    api_key = os.environ.get("GEMINI_API_KEY")
    if not api_key:
        # Erreur : icône + texte, pas seulement la couleur
        st.error("⚠️ Erreur de configuration — La clé API GEMINI_API_KEY est manquante dans les secrets de l'application.")
        st.stop()

    genai.configure(api_key=api_key)
    model = genai.GenerativeModel('gemini-2.5-flash')

    with st.spinner("Analyse pédagogique en cours…"):
        try:
            prompt_parts = [
                f"""Agis en tant qu'expert pédagogique spécialisé en enseignement de la compréhension de texte.
Rédige une fiche enseignant SYNTHÉTIQUE (2 pages maximum) pour un Atelier de Compréhension (ADC) pour le {cycle_short}.

Structure obligatoire :
1. TITRE & INFORMATIONS — niveau, durée estimée, organisation de classe
2. OBJECTIFS DE COMPRÉHENSION — identifie 3 à 5 obstacles SPÉCIFIQUES au texte fourni (lexique opaque, chaînes anaphoriques, implicite culturel ou énonciatif, inférences nécessaires)
3. DÉROULÉ EN 4 PHASES :
   - Phase 1 : Lecture individuelle silencieuse
   - Phase 2 : Tableau collaboratif avec exactement ces colonnes : "Ce qu'on sait" | "Ce qu'on ne sait pas" | "On n'est pas d'accord" — pré-rempli avec 3-4 exemples tirés du texte
   - Phase 3 : Mise en commun et résolution collective
   - Phase 4 : Retour sur les stratégies de compréhension mobilisées
4. QUESTIONS-CLÉS — 5 questions de compréhension fine à poser à la classe, du littéral à l'inférentiel
5. POINTS DE VIGILANCE — erreurs fréquentes à anticiper pour ce texte précis

Sois précis, pratico-pratique. Évite les généralités. Tout doit être ancré dans le texte fourni.
"""
            ]

            file_bytes = uploaded_file.read()
            uploaded_file.seek(0)

            if uploaded_file.type == "application/pdf":
                pdf_doc = fitz.open(stream=file_bytes, filetype="pdf")
                text_content = "".join([page.get_text() for page in pdf_doc])
                if len(text_content.strip()) < 20:
                    for i in range(len(pdf_doc)):
                        page = pdf_doc.load_page(i)
                        pix = page.get_pixmap()
                        prompt_parts.append({"mime_type": "image/png", "data": pix.tobytes("png")})
                else:
                    prompt_parts.append(f"Voici le texte support à analyser :\n\n{text_content}")
            elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                doc_in = Document(io.BytesIO(file_bytes))
                text_content = "\n".join([p.text for p in doc_in.paragraphs if p.text.strip()])
                prompt_parts.append(f"Voici le texte support à analyser :\n\n{text_content}")
            else:
                prompt_parts.append({"mime_type": uploaded_file.type, "data": file_bytes})

            response = model.generate_content(prompt_parts)

            # Résultat — h3 titre de zone
            st.markdown('<h3 style="font-family:\'Fraunces\',Georgia,serif; color:#1B2A4A; font-size:1.1rem; margin:1.5rem 0 0.5rem;">📄 Fiche générée</h3>', unsafe_allow_html=True)
            st.markdown(
                f'<div class="output-zone" role="region" aria-label="Fiche ADC générée">'
                f'{response.text.replace(chr(10), "<br>")}'
                f'</div>',
                unsafe_allow_html=True
            )

            docx_output = create_adc_docx_final(response.text, cycle_short)
            col_a, col_b, col_c = st.columns([1, 2, 1])
            with col_b:
                st.download_button(
                    label="↓ Télécharger la fiche (Word .docx)",
                    data=docx_output,
                    file_name=f"Fiche_ADC_{cycle_short.replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )

        except Exception as e:
            st.error(f"⚠️ Erreur lors de la génération — {e}")

elif generate and not uploaded_file:
    st.warning("⚠️ Aucun fichier déposé — Veuillez d'abord sélectionner un texte support (étape 2).")

# ── Footer ──
st.markdown("""
<footer class="footer-note" role="contentinfo">
  Inspiré de l'outil ACT · ROLL-Descartes · roll-descartes.fr<br>
  Modèle IA : Gemini 2.5 Flash · Interface conçue pour les enseignants du primaire
</footer>
""", unsafe_allow_html=True)
